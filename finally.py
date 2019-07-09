from pprint import pprint

import requests
from docx import Document

from scrapy.http import HtmlResponse

course_dic={"药物化学":133816,
            "药理学":133815,
            "毛中特":133471,
            "生药学":133780,
            "临床医学概论":133763,
            "天然药物化学":98799,
            "15级药化":71676,
            "药用高分子材料学":132904,
            "药事管理学":98789
            }
pprint(course_dic)
course_id=int(input(">>>输入课程id:"))
course=input(">>>输入课程名:")
headers={"Cookie":"PHPSESSID=289ffef8-b982-4eae-bde1-cffc64157449", "User-Agent":"okhttp/3.11.0"}
data={"page":0, "size":30, "courseId":course_id}
#用requests获取题目链接
try:
    urls=[i['link'] for i in requests.post('http://39.105.253.51:90//app/courseTest/list/history',headers=headers,data=data).json()['data']['list']]
except Exception:
    print(requests.post('http://39.105.253.51:90//app/courseTest/list/history',headers=headers,data=data).json())
print(len(urls))
pprint(urls)
def parse(url):
    #处理每个链接网址
    response=HtmlResponse(url=url,body=requests.get(url,headers=headers).content.decode('utf8').encode('utf8'))
    #用xpath和css提取到所需内容
    title = (i.strip() for i in response.css('span.item-title::text').extract() if i.strip())
    choice_before=(i.strip() for i in response.css('div.item-media::text').extract() if i.strip())
    choice_after=[i.strip() for i in response.css('div.item-title::text').extract()]
    answer=(i for i in response.xpath("//span[@style='color: #00B83F;']/text()").extract())
    choice=(j+choice_after[i] for i,j in enumerate(choice_before))
    result=[]
    try:
        result.append(title.__next__())
        result.append(choice.__next__())
        for i in choice:
            result.append(i)
            if i[0]=="A":
                result.insert(-1,answer.__next__())
                result.insert(-1,title.__next__())
        result.append(answer.__next__())
    except StopIteration:
        print("Iter error happend")
    return result
#得到所有题目的列表，并将列表写入文件
res=[]
for i in urls:
    res +=parse(i)
pprint(res)
iter_res = iter(res)
doc = Document()
doc.add_heading(next(iter_res))
for line in iter_res:
    if line[0] == "正":
        doc.add_paragraph(line)
        try:
            doc.add_heading(next(iter_res))
        except Exception:
            pass
    else:
        doc.add_paragraph(line)
doc.save(course+'.docx')
#  f.writelines(res)
